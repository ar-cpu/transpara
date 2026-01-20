#!/usr/bin/env python3
"""
Script to create test PDF and DOCX files with political text examples for Transpara testing
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Test texts organized by political leaning
test_texts = {
    "LEFT-LEANING EXAMPLES": [
        {
            "title": "Universal Healthcare",
            "text": "We need universal healthcare to ensure everyone can access quality medical treatment regardless of their income. The current system leaves millions without coverage while insurance companies profit from human suffering. Affordable housing and living wages are fundamental rights that should be guaranteed to all working families."
        },
        {
            "title": "Social Justice and Equality",
            "text": "Systemic inequality persists because the wealthy elite control the means of production and exploit workers. We must invest in education, protect union rights, and ensure environmental justice for communities that have been historically marginalized. Progressive policies that support equity and inclusion are essential for a just society."
        },
        {
            "title": "Climate Action",
            "text": "Climate change is the greatest threat facing humanity. We must immediately transition to renewable energy, invest in green infrastructure, and hold corporations accountable for their environmental impact. The future of our planet depends on bold action now, not incremental changes that protect corporate profits."
        },
        {
            "title": "Workers' Rights",
            "text": "Workers deserve living wages, strong union protections, and safe working conditions. The gap between CEO pay and worker wages is obscene and must be addressed through progressive taxation and worker ownership. Collective bargaining is a fundamental right that empowers employees to negotiate fair compensation."
        }
    ],
    "CENTER-LEANING EXAMPLES": [
        {
            "title": "Data-Driven Policy",
            "text": "According to recent economic data, unemployment stands at 3.7 percent while GDP growth reached 2.5 percent this quarter. Research suggests that both market forces and targeted government programs play important roles. Studies indicate that a balanced approach combining fiscal responsibility with social investment yields the best outcomes for communities."
        },
        {
            "title": "Moderate Solutions",
            "text": "Evidence-based policy requires careful analysis of multiple factors. The data shows mixed results across different regions, suggesting that one-size-fits-all solutions are ineffective. A pragmatic approach that considers both economic efficiency and social needs appears most sustainable according to recent studies."
        },
        {
            "title": "Bipartisan Cooperation",
            "text": "Effective governance requires compromise and finding common ground. Both parties have valid concerns about healthcare costs, education quality, and economic growth. The best solutions often emerge from combining ideas from different perspectives rather than rigid adherence to ideology."
        },
        {
            "title": "Economic Analysis",
            "text": "Economic indicators show a complex picture. While inflation has moderated, interest rates remain elevated. Consumer spending patterns suggest cautious optimism. Analysts report that multiple factors influence market performance, making simple explanations inadequate."
        }
    ],
    "RIGHT-LEANING EXAMPLES": [
        {
            "title": "Conservative Values",
            "text": "Traditional values and constitutional rights form the foundation of American liberty. Free markets drive innovation and create prosperity when government stays out of the way. Lower taxes, fiscal responsibility, and individual freedom are essential principles that have built this great nation. Law and order must be maintained to protect our communities."
        },
        {
            "title": "Limited Government",
            "text": "The Constitution protects individual rights against government overreach. Free enterprise and competition create jobs and economic growth far better than central planning. Border security is national security, and we must protect American sovereignty. Religious liberty and family values are fundamental to our society's strength."
        },
        {
            "title": "Free Market Principles",
            "text": "Capitalism rewards hard work and innovation. Excessive regulation kills jobs and crushes small business owners. Lower taxes stimulate economic growth and allow families to keep more of what they earn. Government dependency weakens the moral fiber of citizens and creates a cycle of poverty."
        },
        {
            "title": "Constitutional Rights",
            "text": "The Second Amendment protects the right to bear arms, which shall not be infringed. Constitutional originalism preserves the founders' original intent. Individual liberty always trumps collective government mandates. The federal government has expanded far beyond its constitutional limits."
        }
    ]
}

def create_docx_file():
    """Create a DOCX file with test examples"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Transpara Test Document - Political Text Examples', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction
    intro = doc.add_paragraph()
    intro.add_run('This document contains various political text examples for testing the Transpara bias detection system. ').bold = True
    intro.add_run('Each section includes examples from different political perspectives: Left-leaning, Center-leaning, and Right-leaning content.')
    
    doc.add_paragraph()  # Spacing
    
    # Add content for each category
    for category, examples in test_texts.items():
        # Category heading
        heading = doc.add_heading(category, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add each example
        for i, example in enumerate(examples, 1):
            # Example title
            example_title = doc.add_heading(f"Example {i}: {example['title']}", level=2)
            
            # Example text
            para = doc.add_paragraph(example['text'])
            para.style = 'Normal'
            
            doc.add_paragraph()  # Spacing between examples
    
    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run('Note: ').bold = True
    footer.add_run('These examples are for testing purposes only. The bias detection results may vary based on the model\'s training data and interpretation.')
    
    # Save the document
    output_path = 'test_examples.docx'
    doc.save(output_path)
    print(f"✓ Created {output_path}")
    return output_path

def create_pdf_file():
    """Create a PDF file with test examples using reportlab"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        output_path = 'test_examples.pdf'
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#1a1a1a',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor='#2c3e50',
            spaceAfter=10,
            spaceBefore=12
        )
        
        subheading_style = ParagraphStyle(
            'CustomSubHeading',
            parent=styles['Heading2'],
            fontSize=12,
            textColor='#34495e',
            spaceAfter=8,
            spaceBefore=10
        )
        
        # Title
        story.append(Paragraph('Transpara Test Document', title_style))
        story.append(Paragraph('Political Text Examples', title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Introduction
        intro_text = ('<b>Introduction:</b> This document contains various political text examples for testing '
                     'the Transpara bias detection system. Each section includes examples from different '
                     'political perspectives: Left-leaning, Center-leaning, and Right-leaning content.')
        story.append(Paragraph(intro_text, styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Add content for each category
        for category, examples in test_texts.items():
            # Category heading
            story.append(Paragraph(category, heading_style))
            story.append(Spacer(1, 0.1*inch))
            
            # Add each example
            for i, example in enumerate(examples, 1):
                # Example title
                example_title = f"Example {i}: {example['title']}"
                story.append(Paragraph(example_title, subheading_style))
                
                # Example text
                story.append(Paragraph(example['text'], styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
            
            story.append(PageBreak())
        
        # Footer note
        story.append(Spacer(1, 0.2*inch))
        note_text = ('<b>Note:</b> These examples are for testing purposes only. The bias detection results '
                    'may vary based on the model\'s training data and interpretation.')
        story.append(Paragraph(note_text, styles['Normal']))
        
        # Build PDF
        doc.build(story)
        print(f"✓ Created {output_path}")
        return output_path
        
    except ImportError:
        # Fallback: Create a simple text file that can be converted to PDF
        print("reportlab not available, creating text file instead...")
        return create_text_file()

def create_text_file():
    """Create a simple text file as fallback"""
    output_path = 'test_examples.txt'
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("TRANSPARA TEST DOCUMENT - POLITICAL TEXT EXAMPLES\n")
        f.write("=" * 60 + "\n\n")
        f.write("This document contains various political text examples for testing.\n\n")
        
        for category, examples in test_texts.items():
            f.write(f"\n{category}\n")
            f.write("-" * 60 + "\n\n")
            
            for i, example in enumerate(examples, 1):
                f.write(f"Example {i}: {example['title']}\n")
                f.write(f"{example['text']}\n\n")
    
    print(f"✓ Created {output_path}")
    return output_path

if __name__ == "__main__":
    print("Creating test documents for Transpara...")
    print()
    
    # Create DOCX
    docx_path = create_docx_file()
    
    # Create PDF
    pdf_path = create_pdf_file()
    
    print()
    print("=" * 60)
    print("Test documents created successfully!")
    print("=" * 60)
    print(f"DOCX file: {os.path.abspath(docx_path)}")
    print(f"PDF file: {os.path.abspath(pdf_path)}")
    print()
    print("You can now upload these files to Transpara for testing!")
